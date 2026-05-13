import docx

replacements = {
    "Integrated within the Python backend utilizing the ReportLab library, this dedicated service dynamically constructs standardized, branded academic transcripts on the fly by formatting raw JSON data into secure, downloadable PDF documents.": "Integrated within the React frontend utilizing the browser's native print rendering engine, this dedicated service dynamically constructs standardized, branded academic transcripts on the fly by formatting JSON data from the API into secure, downloadable PDF documents via CSS print media queries.",
    
    "c. Pass aggregated data to the Document Generation Layer (ReportLab). d. Render official PDF transcript containing university headers, semester breakdowns, and final classification. e. Securely download PDF to Admin's local machine.": "c. Pass aggregated data to the Document Generation Layer (React Frontend). d. Render official PDF transcript containing university headers, semester breakdowns, and final classification using browser-native print rendering. e. Securely save PDF to Admin's local machine.",
    
    "ReportLab: A powerful software library utilized within the Python ecosystem to programmatically generate secure, multi-page PDF documents, serving as the core engine for transcript rendering.": "Browser Native Print Engine: A powerful rendering mechanism utilized within the React ecosystem to programmatically generate secure, multi-page PDF documents via CSS print media queries, serving as the core engine for transcript rendering.",
    
    "ReportLab was selected over basic HTML-to-PDF converters because it directly manipulates the PDF canvas, ensuring documents cannot be easily tampered with post-generation.": "The native browser print engine combined with React was selected because it allows for pixel-perfect dynamic rendering of transcripts without putting heavy processing load on the backend server, ensuring instant document generation.",
    
    "The integration of the ReportLab PDF engine provided a seamless path for escalating digital records into printable, official university documents.": "The integration of the React-based native print engine provided a seamless path for escalating digital records into printable, official university documents."
}

def replace_text_in_docx(doc_path, output_path):
    print(f"Loading {doc_path}...")
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error loading document: {e}")
        return

    print("Document loaded. Searching for text...")
    replaced_count = 0

    # Iterate through paragraphs
    for p in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in p.text:
                # Basic replacement: this strips run-level formatting (bold/italic) for the replaced string
                # but is the safest way to replace across runs
                p.text = p.text.replace(old_text, new_text)
                replaced_count += 1
                print(f"Replaced an instance in a paragraph.")

    # Iterate through tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)
                            replaced_count += 1
                            print(f"Replaced an instance in a table.")

    if replaced_count > 0:
        print(f"Saving to {output_path}...")
        doc.save(output_path)
        print("Success! Made replacements and saved new document.")
    else:
        print("No matches found. Ensure the strings exactly match.")

if __name__ == "__main__":
    replace_text_in_docx("BLOWSOMED RISING LEGACY PROJECT REPORT.docx", "BLOWSOMED RISING LEGACY PROJECT REPORT - UPDATED.docx")
