from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def add_heading(text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None

def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

# TITLE PAGE
doc.add_paragraph("\n\n")
add_heading("RISING LEGACY: AN AUTOMATED ELECTRONIC GRADING AND TRANSCRIPT MANAGEMENT SYSTEM", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph("\n\n\n\n")
add_paragraph("BY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
doc.add_paragraph("\n\n")
add_paragraph("DEREK-AYEMERE RHEMA OSEGODUWA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_paragraph("SCN/CSC/220880", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
doc.add_paragraph("\n\n\n\n")
add_paragraph("A PROJECT WORK\nPRESENTED TO THE DEPARTMENT OF COMPUTER SCIENCE, FACULTY OF COMPUTING, IN PARTIAL FULFILMENT OF THE REQUIREMENTS FOR THE AWARD OF BACHELOR DEGREE OF SCIENCE (B.Sc.) IN COMPUTER SCIENCE OF BENSON IDAHOSA UNIVERSITY, BENIN CITY, EDO STATE, NIGERIA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
doc.add_paragraph("\n\n\n\n")
add_paragraph("APRIL, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
doc.add_page_break()

# DECLARATION
add_heading("DECLARATION", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("I, Derek-Ayemere Rhema Osegoduwa, hereby declare that this project work was carried out by me under the supervision of Mr. Melvin Omoraro, in the Department of Computer Science, Benson Idahosa University, Benin City, in partial fulfilment of the requirements for the award of the Bachelor of Science (B.Sc.) degree as prescribed by the University.")
doc.add_paragraph("\n\n_______________________________                        ____________________________\nDEREK-AYEMERE RHEMA OSEGODUWA                                  DATE\n   (STUDENT)")
doc.add_page_break()

# CERTIFICATION
add_heading("CERTIFICATION", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("This is to certify that DEREK-AYEMERE RHEMA OSEGODUWA matriculation number SCN/CSC/220880 carried out this project work. In partial fulfilment of the award of Bachelor Degree (B.Sc.) in Computer Science, BENSON IDAHOSA UNIVERSITY undersigned by the following people:")
doc.add_paragraph("\n\n_______________________________                        _________________________\nDEREK-AYEMERE RHEMA OSEGODUWA                               DATE\n (STUDENT)")
doc.add_paragraph("\n\n_______________________________                        ________________________\nMR. MELVIN OMORARO                                                               DATE\n(PROJECT SUPERVISOR)")
doc.add_page_break()

# DEDICATION
add_heading("DEDICATION", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("This work is dedicated to Almighty God, whose infinite wisdom and grace has been my constant guide throughout this journey. His love, protection and blessings is the only reason for the completion of this project possible. To God be all the glory.")
doc.add_page_break()

# ACKNOWLEDGMENTS
add_heading("ACKNOWLEDGMENTS", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("I wish to express my profound gratitude to the Almighty GOD for bestowing upon me the gift of life, wisdom, and grace, which enabled me to successfully complete this seminar report.")
add_paragraph("I am deeply grateful to my supervisor, Mr. Melvin Omoraro for his leadership and guidance.")
add_paragraph("I extended my sincere appreciation to the Head of Department, Dr Sunday Agu and my dedicated lecturers Prof K.O Obahiagbon, Mrs Iriagbonse A. Inyang, Mrs. Cynthia Orie, and Mr. Oyeyemi, for their invaluable contributions to my academic development.")
add_paragraph("My heartfelt thanks goes to my father Mr Mukoro and my siblings for their unwavering love, care and support throughout this journey.")
add_paragraph("I also sincerely appreciate Alvin Ogboru, for his assistance with the design and drafting of this project. I further extend my gratitude to my friends and colleagues who contributed their time, support, and effort towards the development of this work, especially in testing and improving the project.")
doc.add_page_break()

# ABSTRACT
add_heading("ABSTRACT", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("The combination of overcrowded academic offices and the high infrastructure costs associated with enterprise university management systems creates significant barriers for localized departments seeking efficient result computation in developing regions such as Benin City, Edo State. Existing literature highlights a critical deficiency: while digital educational solutions continue to expand, there remains a severe lack of secure, lightweight web platforms explicitly optimized for departmental-level academic triage and low-bandwidth campus networks. Consequently, lecturers and examination officers frequently resort to unsecure physical ledgers or fragmented spreadsheet applications that lack strict academic data compliance and structured digital archiving. To bridge this administrative and technological gap, this study details the development of the Rising Legacy System, a secure, cloud-native automated grading and transcript management system designed to facilitate accurate, real-time academic record keeping.")
add_paragraph("The project utilized the Agile Software Development Life Cycle (SDLC) to engineer a responsive cross-platform solution using React and Tailwind CSS, supported by a highly robust Python FastAPI backend. Instantaneous mathematical grading and automated grade-point conversion were achieved using asynchronous Python logic, alongside dynamic PDF generation to ensure official transcripts can be exported immediately. Data integrity and academic privacy were guaranteed through JSON Web Tokens (JWT) for Role-Based Access Control (RBAC), and an ACID-compliant distributed TiDB database architecture. Furthermore, deployment optimization utilizing Vite successfully reduced the application's client-side footprint, ensuring high accessibility for university staff with limited network capacity. Ultimately, the Rising Legacy platform delivers a highly scalable, legally compliant blueprint for academic administration that successfully decongests physical departmental offices and democratizes verified transcript access.")
doc.add_page_break()

# TABLE OF CONTENTS
add_heading("TABLE OF CONTENTS", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("Title Page\t\ti\nDeclaration\t\tii\nCertification\t\tiii\nDedication\t\tiv\nAcknowledgements\t\tv\nAbstract\t\tvi\nTable of Contents\t\tvii\nList of Figures\t\tviii\nList of Tables\t\tix")
add_paragraph("\nCHAPTER ONE: INTRODUCTION\n1.1 Background to the Study\n1.2 Statement of the Problem\n1.3 Objectives of the Study\n1.4 Scope of the Study\n1.5 Significance of the Study\n1.6 Limitations of the Study\n1.7 Definition of Terms")
add_paragraph("\nCHAPTER TWO: LITERATURE REVIEW\n2.1 Introduction\n2.2 Related Works\n2.3 Summary of Related Works\n2.4 Research Gap")
add_paragraph("\nCHAPTER THREE: METHODOLOGY\n3.1 Adopted Methodology\n3.2 Analysis of the Existing System\n3.2.1 Limitations or Drawbacks of the Existing System\n3.3 Analysis of the Proposed System\n3.3.1 Preliminary Design\n3.3.2 Proposed System Justification\n3.3.3 Benefits of the Proposed System\n3.4 System Design\n3.5 Database Design\n3.6 Input/Output Specification\n3.7 Cost Analysis")
add_paragraph("\nCHAPTER FOUR: IMPLEMENTATION AND DISCUSSION\n4.1 Programming Language of Implementation\n4.1.1 Justification of Programming Languages Used\n4.2 System Requirements\n4.3 Implementation Guidelines\n4.4 Results and Discussion")
add_paragraph("\nCHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATIONS\n5.1 Summary of Major Findings\n5.2 Conclusion\n5.3 Recommendations\n\nREFERENCES")
doc.add_page_break()

doc.save("BLOWSOMED RISING LEGACY PROJECT REPORT - REWRITTEN.docx")
print("Initial chapters created successfully.")
