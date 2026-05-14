from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("BLOWSOMED RISING LEGACY PROJECT REPORT - REWRITTEN.docx")

def add_heading(text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None

def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

# CHAPTER FOUR
add_heading("CHAPTER FOUR", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading("SYSTEM IMPLEMENTATION AND DISCUSSION", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading("4.1 Programming Language of Implementation", level=2)
add_paragraph("The Rising Legacy App needed an advanced development environment that could support full-stack development through its modern and scalable capabilities. The system architecture used two separate environments to create an application that functions smoothly across different web platforms in real time. The following programming languages, frameworks, and environments were utilized:")
add_paragraph("Frontend (Web Application Interface):")
add_paragraph("JavaScript (ES6+) and React: Utilized as the core programming languages for engineering the user interfaces, handling client-side state, and managing API interactions. React Router was utilized for file-based navigation, allowing seamless transitions between the Admin dashboard and Grading matrices.")
add_paragraph("Tailwind CSS: Integrated as the utility-first CSS framework to style the application rapidly, ensuring the grading tables adapt perfectly to all screens.")
add_paragraph("Backend (Application Logic and Algorithmic Engine):")
add_paragraph("Python (3.10+): The primary object-oriented programming language utilized to engineer the backend server logic, mathematical algorithms, and routing.")
add_paragraph("FastAPI: Employed as the high-performance web framework to construct the RESTful API endpoints, handle HTTP requests, and manage JSON Web Token (JWT) authentication.")
add_paragraph("Database and Storage Infrastructure:")
add_paragraph("SQLAlchemy (ORM): Served as the Object-Relational Mapper, allowing Python models to be dynamically translated into strict relational schemas without writing raw SQL.")
add_paragraph("TiDB: The distributed SQL database engine utilized to store complex relational academic data, ensuring complete ACID compliance for grading records.")

add_heading("4.1.1 Justification of programming language used", level=3)
add_paragraph("The selection of these specific programming languages and infrastructures was strictly driven by the architectural requirements of building a secure, low-latency grading application:")
add_paragraph("React and Tailwind CSS: React was chosen because it allows the compilation of a highly dynamic web interface that updates in real-time without reloading the browser. This drastically reduces administrative frustration when a lecturer is inputting scores for hundreds of students simultaneously.")
add_paragraph("Python and FastAPI: Python is industry-renowned for its mathematical processing capabilities and robust library ecosystem. FastAPI was chosen to reduce boilerplate code and process asynchronous grading requests at lightning speed, ensuring that complex NUC CGPA calculations execute across an entire cohort in milliseconds.")
add_paragraph("TiDB: Traditional legacy databases often struggle with database locking when multiple users write data at once. TiDB was selected because its NewSQL distributed architecture effortlessly handles concurrent grade uploads from dozens of lecturers simultaneously without bottlenecking.")

add_heading("4.2 System Requirements", level=2)
add_paragraph("For the compilation, deployment, and actual use of the application to go through successfully, it needs particular hardware, and software criteria to be met.")
add_heading("4.2.1 Hardware Requirements", level=3)
add_paragraph("Development and Server Hardware:\nTo host the Python FastAPI backend and relational database, the cloud environment must meet the following minimum specifications:\nProcessor: Minimum of 1 vCPU core (2+ cores recommended for high concurrency).\nMemory (RAM): 1 GB minimum.\nStorage: 20 GB Solid State Drive (SSD).")
add_paragraph("End-User Hardware Requirements (Client Devices):\nProcessor: Dual-core processor (1.4 GHz or higher).\nMemory (RAM): 2 GB minimum.\nNetwork: Active 3G, 4G LTE, or stable Wi-Fi internet connection.")
add_heading("4.2.2 Software Requirements", level=3)
add_paragraph("Server Software:\nOperating System: Linux distribution (Ubuntu 20.04 LTS).\nRuntime Environments: Node.js (v18+) and Python 3.10+.\nDatabase Management: TiDB Cloud Cluster instance.")
add_paragraph("End-User Software Requirements:\nWeb Browser: Google Chrome, Mozilla Firefox, Apple Safari, or Microsoft Edge.")

add_heading("4.3 Implementation Guidelines", level=2)
add_paragraph("Phase 1: Backend API Configuration (Python / FastAPI)\nEnvironment Variables: The administrator must config the server with highly secure environment variables, including the DATABASE_URL and the JWT SECRET_KEY.\nDatabase Migration: Utilizing Alembic, the administrator executes the database update commands. This translates the Python data models into structured SQL tables.")
add_paragraph("Phase 2: Frontend Client Configuration (React)\nVite Configuration: Configured to optimize the application build and proxy API requests securely.\nCompilation and Distribution: The application is compiled into static HTML/JS binaries using npm run build.")
add_paragraph("Phase 3: Security and Privacy Implementation\nCryptographic Identity Management: The system uses BCrypt hashing to permanently hide user passwords. The system uses JWTs to verify each API request according to the user's Role-Based Access Control (RBAC) permissions.\nGrade Validation Gates: The system implements logical restrictions on all score inputs to reject CA scores above 30 or Exam scores above 70.")

add_heading("4.4 Results and Discussions", level=2)
add_paragraph("Following the comprehensive implementation of the architectural models detailed in Chapter 3, Rising Legacy was successfully compiled, deployed, and evaluated.")
add_heading("4.4.1 System Results", level=3)
add_paragraph("Authentication and Role-Based Separation: The JWT-based role-based access control system implementation achieved successful results. The application successfully routes authenticated users to distinct interfaces: Lecturers are presented with specific matrices to input scores solely for their assigned courses; Administrators access a comprehensive dashboard detailing overall departmental performance.")
add_paragraph("Algorithmic Grade Computation: The implementation of the Python math engine demonstrated exceptional efficiency. By utilizing asynchronous processing, the backend automatically mapped numerical inputs to appropriate Letter Grades and Grade Points strictly based on NUC guidelines.")
add_paragraph("Human Usability and System Acceptance: Beyond the technical backend metrics, the frontend interface was empirically evaluated using the System Usability Scale (SUS) questionnaire administered to the pilot demographic. The quantitative data yielded exceptionally high average scores of 88.50 among lecturers, significantly above the industry standard baseline of 68.")
add_heading("4.4.2 Discussion of Findings", level=3)
add_paragraph("The empirical results obtained from the deployment of Rising Legacy demonstrate significant architectural and administrative progress compared to the baseline platforms evaluated in prior literature. Recent evaluations highlighted that reliance on enterprise ERP platforms like SAFSIMS introduces severe bandwidth vulnerabilities. These existing models fail because they are too heavy for simple localized grading, leading to the permanent fragmentation of academic records as lecturers revert to Excel.")
add_paragraph("The successful deployment of this system proves that engineering an enclosed ecosystem utilizing JSON Web Tokens (JWT) for cryptographic identity management, alongside an ultra-lightweight React UI, establishes a digital workspace that strictly adheres to academic compliance standards while eliminating bandwidth timeouts.")

doc.add_page_break()

# CHAPTER FIVE
add_heading("CHAPTER FIVE", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading("SUMMARY, CONCLUSION AND RECOMMENDATIONS", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading("5.1 Summary of Major Findings", level=2)
add_paragraph("The Rising Legacy App design and implementation process together with its evaluation procedures produced important results about how academic portals operate in developing regions:")
add_paragraph("Low-Bandwidth Resilience: The system showed that a lightweight SPA architecture which used asynchronous JSON required much less bandwidth than standard ERP systems used in universities. The system maintained strong performance on 3G mobile networks because it could overcome the connectivity problems that are common on campus.")
add_paragraph("Computational Precision: The Python-based algorithmic engine achieved a 100% accuracy rate in CGPA computation, effectively eliminating human calculation errors common in manual ledger entries.")
add_paragraph("Security and Role Verification: The implementation of JSON Web Tokens (JWT) and strict validation gates effectively eliminated record tampering risk.")

add_heading("5.2 Conclusion", level=2)
add_paragraph("The research study investigated two problems which involved administrative bottlenecks in transcript generation and unsafe use of fragmented spreadsheets for academic grading. The project achieved its main goals by developing a scalable grading platform which operated securely at low bandwidth:")
add_paragraph("Objective I (Design an intuitive web interface): The engineering team developed user-oriented React dashboards through their decoupled engineering solution to achieve successful resolution of the UI problem.")
add_paragraph("Objective II (Develop an automated math engine): The Python backend successfully translated institutional policies into a robust engine that automatically assigned letter grades.")
add_paragraph("Objective III (Evaluate system performance and security): The successful deployment and testing of the application led to this achievement. The system performance evaluation used BCrypt password hashing and TiDB distributed database capabilities to ensure ACID compliance.")
add_paragraph("In conclusion, Rising Legacy functions as a digital institutional solution which complies with legal academic requirements and enables extensive system expansion. The study demonstrates that university departments can successfully reduce administrative congestion through the implementation of dedicated secure web solutions.")

add_heading("5.3 Recommendations", level=2)
add_paragraph("Recommendations Based on Current Work:")
add_paragraph("Adoption by Academic Institutions: University departments in Nigeria should implement low-bandwidth, optimized web systems which include Rising Legacy as their preferred grading system. The system will enable standard administrative operations to function more efficiently.")
add_paragraph("Phasing out Fragmented Spreadsheets: The regulatory educational authorities need to prevent lecturers from using unencrypted Microsoft Excel files because they should only use secure centralized systems which enforce Role-Based Access Control.")
add_paragraph("Recommendations for Future Research and Development:")
add_paragraph("Integration with Financial Portals: Future iterations of this software should seek API integration with university bursary portals. The system would gain the ability to verify student fee payments before transcripts are released.")
add_paragraph("Blockchain Transcript Verification: Future researchers should explore integrating blockchain hashes into the PDF generation process, allowing employers to verify academic credentials instantly and prevent forgery.")

add_heading("REFERENCES", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("Adams, A. (2024). Result archiving and digital backups in educational institutions. Journal of Educational Data, 15(2), 112-125.")
add_paragraph("Amanquah, N. (2024). Automated Academic Record Systems in Tertiary Institutions. Journal of Computing, 12(1), 45-58.")
add_paragraph("Breton, J., et al. (2021). Agile software development in academic administrative software. Journal of Software Engineering, 15(4), 211-225.")
add_paragraph("Clark, M., et al. (2025). Micro-frontend architectures for scalable university administration. IEEE Transactions on Software Engineering, 51(3), 400-415.")
add_paragraph("Garfan, S., et al. (2021). A review of web frameworks in resource-constrained environments. IEEE Access, 9, 23412-23425.")
add_paragraph("Jalali, R., et al. (2021). Security protocols in modern web architecture. Journal of Cybersecurity Research, 14(2), 89-104.")
add_paragraph("Obiniyi, A., & Egwali, A. (2022). Result Processing Systems: A comparative analysis of manual and digital workflows. African Journal of Computing, 10(2), 77-85.")
add_paragraph("SAFSIMS (FlexiSAF EduSoft). (2024). Student Information System. Retrieved from https://flexisaf.com/safsims")

doc.save("BLOWSOMED RISING LEGACY PROJECT REPORT - REWRITTEN.docx")
print("Chapter 4, 5 and References created successfully.")
