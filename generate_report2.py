from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("BLOWSOMED RISING LEGACY PROJECT REPORT - REWRITTEN.docx")

def add_heading(text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = align
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None

def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

# CHAPTER ONE
add_heading("CHAPTER ONE", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading("INTRODUCTION", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading("1.1 Background to the Study", level=2)
add_paragraph("The management of academic records has depended on manual collation and physical file storage since the inception of formal educational administration. The implementation of Information and Communication Technology (ICT) across all countries created permanent transformations which changed fundamental aspects of institutional operations. Automated grading and record management systems function as the primary disruptive technology which academic administrators use to deliver accurate, decentralized data processing through advanced digital networks (Amanquah, 2024). The educational technology sector provides multiple management services through an extensive service model that includes centralized database processing, cloud-based university information management systems (UIMS), and instant digital archiving systems. Educational institutions are increasingly recognizing these technological developments as essential elements which must be implemented to achieve seamless administrative operations rather than just convenient enhancements. Digital academic solutions provide equal access which enables faculty members to streamline their workflows while significantly reducing the financial and temporal costs associated with physical record keeping.")
add_paragraph("The educational technology industry has experienced exceptional expansion during the past five years which resulted in considerable financial evaluations and the sector will continue to digitize core academic functions at an accelerated rate. The academic sector expands its digital footprint rapidly because three main factors interact with each other: educational facilities are replacing their traditional ledger methods with modern distributed databases, administrators require instant, reliable access to historical student data, and students globally demand verifiable digital transcripts to compete in a fast-paced employment market. The global expansion of secure web frameworks has created vast access to academic management services, acting as the primary solution for departments seeking to eradicate manual computational errors.")
add_paragraph("The tertiary education system in developing nations faces a crisis because of extreme deficiencies in automated processing which results in administrative bottlenecks. Academic departments experience constant workflow congestion because individual lecturers compute grades on disparate spreadsheets, and examination officers must spend weeks harmonizing these files to generate final results. The system functions in an inefficient manner because it causes processing delays which lead to delayed graduations and creates hazardous situations where student records are susceptible to loss or misplacement. The Automated Electronic Grading and Transcript Management System, titled Rising Legacy, will serve as a software solution that enables administrators and lecturers to connect through a centralized web platform. The system uses modern software engineering frameworks which include React with Tailwind CSS for its responsive frontend and Python (FastAPI) with a TiDB distributed database for its secure backend to create a dedicated platform that enables faculty to securely compute results and generate transcripts instantly.")

add_heading("1.2 Statement of the Problem", level=2)
add_paragraph("The worldwide growth of educational technology faces serious challenges locally because academic departments cannot process results efficiently without manual intervention, which results in extended administrative wait times, delayed transcript generation, and creates significant distress for graduating students seeking further education. The existing semi-manual solutions utilized by many departments face major issues because they depend heavily on unencrypted Microsoft Excel spreadsheets and physical flash drives which leave data highly vulnerable to tampering or corruption. Furthermore, the reliance on manual mathematical computation for complex CGPA tracking inherently introduces a high margin of human error. The use of these common, fragmented methods for academic record management creates serious risks which stem from their failure to provide Role-Based Access Control (RBAC), their lack of highly available cloud backups, and their inability to generate instantaneous, standardized transcripts. The fundamental problem requires a dependable, automated grading portal which functions efficiently to synchronize and secure data. The proposed system, Rising Legacy, will function as a complete software engineering solution that will enable academic professionals to conduct accurate result processing through its lightweight, secure architecture, which only authorized faculty staff will use for grading and academic assessments.")

add_heading("1.3 Aim and Objectives of the study", level=2)
add_paragraph("The aim of this study is to develop Rising Legacy, a web-based Automated Electronic Grading and Transcript Management System which enables lecturers and administrators to securely interact, compute academic results, and instantaneously generate transcripts. The specific objectives of this study are to:")
add_paragraph("1. To design an intuitive, responsive web interface for seamless lecturer and administrator interaction during score input and verification.")
add_paragraph("2. To develop a Python-based automated logical engine that accurately computes GPA and CGPA strictly adhering to National Universities Commission (NUC) guidelines.")
add_paragraph("3. To implement a secure backend and distributed TiDB database framework with role-based access control (RBAC) to facilitate reliable, authorized record management.")
add_paragraph("4. To evaluate the usability, performance, and security of the developed system through software testing procedures to ensure that academic records remain protected and transcripts are generated instantly in PDF format.")

add_heading("1.4 Scope of the study", level=2)
add_paragraph("The research focuses on creating and testing Rising Legacy, an Automated Electronic Grading and Transcript Management System which operates smoothly on web browsers across various devices. The research takes place exclusively within the Department of Computer Science at Benson Idahosa University, located in Benin City, Edo State, Nigeria. The study focuses on authenticated lecturers, departmental examination officers, and system administrators who handle continuous assessments and exam grading within this specific academic demographic.")
add_paragraph("The project implements secure user registration together with a role-based access control (RBAC) system, automated grade calculation algorithms, and real-time PDF transcript generation functions. The study strictly limits its coverage to departmental result processing and academic tracking. It excludes university-wide financial integrations, student tuition portals, admissions processing, and integration with legacy national electronic health or general ERP record systems.")

add_heading("1.5 Significance of the study", level=2)
add_paragraph("Research projects exist to solve specific problems while their findings bring progress to a particular academic discipline. The study holds academic significance because it delivers educational benefits and practical advantages while producing solutions for policy-related problems.")
add_paragraph("Academic Contribution: The research demonstrates its contribution to computer science and software engineering through its development of a highly available, cloud-native educational management system. The research establishes a practical method for university departments to implement distributed SQL databases (TiDB) and modern web architectures which perform efficiently in academic environments. The study provides a reference material which future researchers can use to study academic informatics and data synchronization architectures.")
add_paragraph("Practical Contribution: The proposed system allows faculty members to instantly and accurately compute student results without relying on manual ledger calculations. The system helps administrators save countless hours typically wasted physically searching filing cabinets and cuts down the waiting time for graduates needing transcripts. The system enables university staff to use a virtual workspace which helps them schedule their grading securely while protecting student privacy.")
add_paragraph("Policy and Institutional Contribution: The study results help university senates and educational policymakers evaluate distributed cloud systems as practical solutions for enhancing academic data integrity. The systems enable secure digital archiving which aligns with National Universities Commission (NUC) compliance guidelines and prevents the falsification or loss of critical academic records.")

add_heading("1.6 Limitations of the study", level=2)
add_paragraph("This study is subject to several limitations that may affect the generalizability and scope of its findings:")
add_paragraph("Regional Variance: The system is evaluated in a localized departmental pilot setting. Results may not fully apply across all university faculties due to differences in credit load structures, distinct departmental grading policies, and varying levels of digital literacy among staff.")
add_paragraph("Technical Dependencies: The platform requires reliable computer hardware and stable internet. Despite frontend optimizations, poor internet connectivity on campus may temporarily disrupt real-time grade uploads to the cloud database.")
add_paragraph("Operational Adaptation Constraints: Transitioning from traditional manual workflows and local spreadsheets to a centralized digital portal requires a behavioral shift. Older faculty members may experience a learning curve when adopting the new digital interface.")
add_paragraph("Evaluation Timeframe: The research is limited by a short academic timeframe. This restricts the long-term assessment of system performance over multiple graduating cohorts, extended user adoption, and sustained administrative impact.")

add_heading("1.7 Definition of terms", level=2)
add_paragraph("CGPA (Cumulative Grade Point Average): A mathematical representation of a student's overall academic performance across multiple semesters, used to determine final degree classifications.")
add_paragraph("Transcript: An official, comprehensive inventory of the courses taken, grades earned, and degrees awarded to a student throughout their entire course of study.")
add_paragraph("RBAC (Role-Based Access Control): A security model that restricts system access based on the role assigned to a user, such as Lecturer, Examination Officer, or System Administrator.")
add_paragraph("FastAPI: A modern, high-performance web framework for building backend APIs with Python, utilized in this system to process grading logic and mathematical algorithms instantly.")
add_paragraph("TiDB: An open-source, distributed SQL database that guarantees ACID compliance and ensures academic records are safely stored with horizontal scalability to prevent record duplication or data loss.")
add_paragraph("JWT (JSON Web Token): A cryptographically secure, compact token used by the backend system to verify the identity and role of a user each time they attempt to access protected grading data.")
add_paragraph("SDLC (Software Development Life Cycle): A structured sequence of stages in software engineering to plan, create, test, and deploy the proposed academic system.")
doc.add_page_break()

doc.save("BLOWSOMED RISING LEGACY PROJECT REPORT - REWRITTEN.docx")
print("Chapter 1 created successfully.")
